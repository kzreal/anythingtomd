"""查找 rId5 图片被引用的位置"""
from docx import Document
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

doc = Document('/Users/kyle/Downloads/测试.docx')

# 查找 rId5 在文档中的所有引用位置
print("🔍 查找 rId5 在文档中的所有引用...\n")

# 检查所有段落
for i, para in enumerate(doc.paragraphs, 1):
    xml_str = str(para._element)
    if 'rId5' in xml_str:
        print(f"\n✅ 段落 {i} 包含 rId5:")
        print(f"   文本: {para.text[:100]}")
        print(f"   XML 片段 (前500字符): {xml_str[:500]}")
        print()

# 检查所有表格
print("\n" + "="*60)
print("检查表格中的 rId5 引用...")
print("="*60 + "\n")

for i, table in enumerate(doc.tables, 1):
    xml_str = str(table._element)
    if 'rId5' in xml_str:
        print(f"\n✅ 表格 {i} 包含 rId5")

        # 检查表格的每个单元格
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_xml = str(cell._element)
                if 'rId5' in cell_xml:
                    print(f"   在单元格 [{row_idx}][{cell_idx}] 中")
                    print(f"   文本: {cell.text[:100]}")

# 检查文档中的所有 drawing 元素
print("\n" + "="*60)
print("检查所有 drawing 元素...")
print("="*60 + "\n")

# 获取文档的 body 部分
body = doc._element.body

# 查找所有包含 rId5 的元素
for elem in body.iter():
    elem_xml = str(elem)
    if 'rId5' in elem_xml:
        print(f"✅ 找到包含 rId5 的元素:")
        print(f"   标签名: {elem.tag}")
        print(f"   完整 XML: {elem_xml[:800]}")
        print()
