"""
使用 Playwright 测试 Word 转 Markdown 功能
"""
import os
from pathlib import Path
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


def create_test_word_with_images():
    """创建包含图片的测试 Word 文档"""
    upload_dir = Path("/tmp/test_word")
    upload_dir.mkdir(exist_ok=True)

    # 创建测试图片
    img_dir = upload_dir / "images"
    img_dir.mkdir(exist_ok=True)

    # 图片 1: 流程图
    img1 = Image.new('RGB', (600, 300), color='#f0f0f0')
    draw = ImageDraw.Draw(img1)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 28)
    except:
        font = ImageFont.load_default()

    # 画流程框
    draw.rectangle([50, 80, 150, 220], fill='#4CAF50', outline='black', width=2)
    draw.rectangle([225, 80, 325, 220], fill='#2196F3', outline='black', width=2)
    draw.rectangle([400, 80, 500, 220], fill='#FF9800', outline='black', width=2)

    # 画箭头
    draw.line([(150, 150), (225, 150)], fill='black', width=3)
    draw.polygon([(220, 145), (225, 155), (220, 155)], fill='black')
    draw.line([(325, 150), (400, 150)], fill='black', width=3)
    draw.polygon([(395, 145), (400, 155), (395, 155)], fill='black')

    # 文字
    draw.text((65, 145), "开始", fill='white', font=font)
    draw.text((240, 145), "处理", fill='white', font=font)
    draw.text((415, 145), "完成", fill='white', font=font)

    img1_path = img_dir / "flowchart.png"
    img1.save(img1_path)
    print(f"✓ 创建测试图片 1: {img1_path}")

    # 图片 2: 架构图
    img2 = Image.new('RGB', (600, 350), color='white')
    draw2 = ImageDraw.Draw(img2)
    try:
        font2 = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 24)
    except:
        font2 = ImageFont.load_default()

    # 画客户端
    draw2.rectangle([200, 30, 400, 80], fill='#E3F2FD', outline='black', width=2)
    draw2.text((260, 45), "客户端应用", fill='black', font=font2)

    # 画服务器
    draw2.rectangle([150, 150, 450, 200], fill='#FFF3E0', outline='black', width=2)
    draw2.text((250, 165), "服务器 API", fill='black', font=font2)

    # 画数据库
    draw2.rectangle([200, 280, 400, 330], fill='#E8F5E9', outline='black', width=2)
    draw2.text((260, 295), "数据库", fill='black', font=font2)

    # 画连接线
    draw2.line([(300, 80), (300, 150)], fill='blue', width=2)
    draw2.line([(300, 200), (300, 280)], fill='blue', width=2)

    img2_path = img_dir / "architecture.png"
    img2.save(img2_path)
    print(f"✓ 创建测试图片 2: {img2_path}")

    # 创建 Word 文档
    doc = Document()

    # 添加标题
    title = doc.add_heading('测试文档 - Word转Markdown', level=1)

    # 添加介绍
    doc.add_paragraph('这是一个测试文档，用于验证 Word 转 Markdown 功能，包括图片识别。')

    doc.add_heading('章节一：流程说明', level=2)
    doc.add_paragraph('以下是业务流程图：')

    # 添加图片 1
    doc.add_picture(str(img1_path), width=Inches(5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('上图展示了完整的业务流程。')

    doc.add_heading('章节二：系统架构', level=2)
    doc.add_paragraph('系统架构设计如下：')

    # 添加图片 2
    doc.add_picture(str(img2_path), width=Inches(5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('系统采用三层架构设计，包括客户端、服务器和数据库层。')

    doc.add_heading('章节三：技术要点', level=2)
    doc.add_paragraph('关键技术点包括：')
    doc.add_paragraph('1. RESTful API 设计', style='List Number')
    doc.add_paragraph('2. 数据库优化', style='List Number')
    doc.add_paragraph('3. 前端响应式设计', style='List Number')

    # 添加表格
    doc.add_heading('章节四：参数配置', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '参数名'
    table.rows[0].cells[1].text = '默认值'
    table.rows[0].cells[2].text = '说明'

    table.rows[1].cells[0].text = 'timeout'
    table.rows[1].cells[1].text = '30'
    table.rows[1].cells[2].text = '超时时间（秒）'

    table.rows[2].cells[0].text = 'retries'
    table.rows[2].cells[1].text = '3'
    table.rows[2].cells[2].text = '重试次数'

    table.rows[3].cells[0].text = 'max_size'
    table.rows[3].cells[1].text = '1024'
    table.rows[3].cells[2].text = '最大大小（MB）'

    # 保存文档
    doc_path = upload_dir / "test_with_images.docx"
    doc.save(doc_path)
    print(f"✓ 创建测试 Word 文档: {doc_path}")

    return doc_path


def test_word_conversion():
    """测试 Word 转 Markdown 功能"""
    print("🧪 开始测试 Word 转 Markdown 功能...\n")

    # 创建测试文档
    test_doc = create_test_word_with_images()
    print(f"\n📄 使用测试文档: {test_doc}\n")

    with sync_playwright() as p:
        # 启动浏览器
        browser = p.chromium.launch(headless=False)
        page = browser.new_page()

        try:
            # 访问应用
            app_url = "http://localhost:8080"
            print(f"🌐 访问应用: {app_url}")
            page.goto(app_url, wait_until='networkidle')
            print("✓ 页面加载成功\n")

            # 上传 Word 文件
            print("📤 上传 Word 文件...")
            file_input = page.locator('input[type="file"]')
            file_input.set_input_files(str(test_doc))
            print("✓ 文件已选择\n")

            # 等待文件分析完成
            print("⏳ 等待文件分析...")
            page.wait_for_selector('#fileInfo', timeout=30000)
            file_info = page.locator('#fileInfo').text_content()
            print(f"✓ 文件分析完成\n")
            print(f"📋 文件信息:\n{file_info}\n")

            # 点击转换按钮
            print("▶️ 开始转换...")
            convert_btn = page.locator('#convertBtn')
            convert_btn.click()
            print("✓ 转换已启动\n")

            # 等待转换完成 - 增加超时时间（图片识别需要更长时间）
            print("⏳ 等待转换完成（包含图片识别）...")
            try:
                page.wait_for_selector('#previewText', timeout=600000)  # 10分钟超时
                print("✓ 转换完成\n")
            except Exception as e:
                print(f"⚠ 等待预览内容时出错: {e}")
                # 尝试检查是否有错误信息
                error_elem = page.locator('.error').first
                if error_elem.count() > 0:
                    error_text = error_elem.text_content()
                    print(f"❌ 页面错误: {error_text}")
                raise


            # 等待额外时间让内容完全加载
            page.wait_for_timeout(2000)

            # 直接获取预览内容（通过 DOM）
            preview_content = page.evaluate("() => document.getElementById('previewText').textContent")
            print(f"📝 预览内容长度: {len(preview_content)} 字符\n")
            print(f"📝 预览内容 (前1500字符):\n{preview_content[:1500]}...\n")

            # 检查是否包含行号
            if '<!--' in preview_content:
                print("✓ 检测到行号\n")
            else:
                print("⚠ 未检测到行号\n")

            # 检查图片识别结果
            if '[图片:' in preview_content:
                print("✓ 检测到图片识别结果")
                # 提取所有图片描述
                import re
                image_descriptions = re.findall(r'\[图片: (.*?)\]', preview_content)
                print(f"📷 识别到的图片描述:")
                for i, desc in enumerate(image_descriptions, 1):
                    print(f"   {i}. {desc}")
                print()
            else:
                print("⚠ 未检测到图片识别结果\n")

            # 截图保存
            screenshot_path = "/tmp/test_word_preview.png"
            page.screenshot(path=screenshot_path, full_page=True)
            print(f"📸 预览截图已保存: {screenshot_path}\n")

            # 测试切片列表
            sections_list = page.locator('#sectionList .section-item').all()
            print(f"📑 切片数量: {len(sections_list)}\n")

            if len(sections_list) > 0:
                # 显示所有切片标题
                print("📋 切片列表:")
                for i, section in enumerate(sections_list):
                    title = section.text_content()
                    print(f"   {i+1}. {title}")
                print()

                # 尝试找到并切换到包含图片的切片
                print("🔍 查找包含图片的切片...")
                for i, section in enumerate(sections_list):
                    section.click()
                    page.wait_for_timeout(500)
                    current_preview = page.locator('#previewText').text_content()

                    if '[图片:' in current_preview:
                        print(f"✓ 找到包含图片的切片: {section.text_content()}")
                        page.wait_for_timeout(500)
                        break
                    else:
                        print(f"   切片 {i+1} 不包含图片")

                # 获取当前预览内容
                preview_content = page.locator('#previewText').text_content()

            if len(sections_list) > 0:
                # 显示所有切片标题
                print("📋 切片列表:")
                for section in sections_list:
                    title = section.text_content()
                    print(f"   - {title}")
                print()

                # 点击第一个切片
                if len(sections_list) > 1:
                    print("🔍 测试切片切换...")
                    first_section = sections_list[0]
                    first_title = first_section.text_content()
                    first_section.click()
                    page.wait_for_timeout(1000)

                    # 获取新预览内容
                    new_preview = page.locator('#previewText').text_content()
                    if new_preview != preview_content:
                        print(f"✓ 已切换到: {first_title}\n")
                    else:
                        print("⚠ 切片内容未变化\n")

                    # 点击第二个切片
                    second_section = sections_list[1]
                    second_title = second_section.text_content()
                    second_section.click()
                    page.wait_for_timeout(1000)
                    print(f"✓ 已切换到: {second_title}\n")

                # 测试下载功能
                print("⬇️ 测试下载功能...")
                with page.expect_download() as download_info:
                    page.locator('#downloadBtn').click()
                download = download_info.value
                print(f"✓ 下载完成: {download.suggested_filename}\n")

                # 保存下载文件
                save_path = upload_dir / download.suggested_filename
                download.save_as(save_path)
                print(f"✓ 文件已保存到: {save_path}\n")

                # 检查 ZIP 文件内容
                import zipfile
                with zipfile.ZipFile(save_path, 'r') as zf:
                    file_list = zf.namelist()
                    print(f"📦 ZIP 文件内容:")
                    for f in sorted(file_list):
                        print(f"   - {f}")
                    print()

                    # 读取并显示第一个 Markdown 文件的内容
                    md_files = [f for f in file_list if f.endswith('.md') and not f.startswith('00_')]
                    if md_files:
                        first_md = md_files[0]
                        with zf.open(first_md) as md_file:
                            md_content = md_file.read().decode('utf-8')
                        print(f"📄 {first_md} 内容 (前500字符):\n{md_content[:500]}...\n")

            print("✅ 所有测试完成")
            return True

        except Exception as e:
            print(f"✗ 测试失败: {type(e).__name__}: {e}")
            import traceback
            traceback.print_exc()

            # 截图保存错误状态
            try:
                screenshot_path = "/tmp/test_word_error.png"
                page.screenshot(path=screenshot_path)
                print(f"📸 错误截图已保存: {screenshot_path}")
            except:
                pass

            return False
        finally:
            browser.close()


if __name__ == "__main__":
    success = test_word_conversion()
    if not success:
        exit(1)
