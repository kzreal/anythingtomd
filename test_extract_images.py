"""测试图片提取函数"""
import sys
sys.path.insert(0, '/Users/kyle/Projects/AlltoMarkdown')

from docx import Document
from converters.word_converter import WordConverter

doc = Document('/Users/kyle/Downloads/测试.docx')
converter = WordConverter('/Users/kyle/Downloads/测试.docx')

# 加载文档
converter.load_document()

# 获取段落 19
para = converter.doc.paragraphs[18]  # 索引从 0 开始

print("="*60)
print("测试段落 19 的图片提取")
print("="*60)
print(f"段落文本: {para.text[:100]}")

# 调用 extract_paragraph_images
import logging
logging.basicConfig(level=logging.DEBUG)
images = converter.extract_paragraph_images(para)

print(f"\n提取到的图片数量: {len(images)}")
for i, img in enumerate(images):
    print(f"  图片 {i+1}:")
    print(f"    ID: {img['id']}")
    print(f"    格式: {img['format']}")
    print(f"    数据大小: {len(img['data']) if img['data'] else 0} 字节")

# 检查文档中所有图片
print("\n" + "="*60)
print("检查文档中所有段落的图片")
print("="*60)

total_images = 0
for i, para in enumerate(converter.doc.paragraphs, 1):
    images = converter.extract_paragraph_images(para)
    if images:
        total_images += len(images)
        print(f"段落 {i}: 找到 {len(images)} 张图片")
        for img in images:
            print(f"  - ID: {img['id']}, 大小: {len(img['data'])} 字节")

print(f"\n总共找到 {total_images} 张图片")
