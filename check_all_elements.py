"""检查文档所有元素中的图片"""
from docx import Document
from docx.oxml.ns import qn
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

doc = Document('/Users/kyle/Downloads/测试.docx')

print("="*60)
print("检查文档所有可能的图片位置")
print("="*60 + "\n")

# 1. 检查段落中的所有 drawing 元素
print("1️⃣ 检查段落中的 drawing 元素...\n")
drawing_count = 0
for i, para in enumerate(doc.paragraphs, 1):
    # 查找所有 w:drawing 元素
    drawings = para._element.xpath('.//w:drawing')
    if drawings:
        drawing_count += len(drawings)
        print(f"✅ 段落 {i} 包含 {len(drawings)} 个 drawing 元素")
        print(f"   文本: {para.text[:100]}")
        for j, drawing in enumerate(drawings):
            print(f"   Drawing {j}: {drawing.tag}")

if drawing_count == 0:
    print("   ⚠️ 段落中未找到 drawing 元素")

# 2. 检查表格中的所有 drawing 元素
print(f"\n2️⃣ 检查表格中的 drawing 元素...\n")
table_drawing_count = 0
for i, table in enumerate(doc.tables, 1):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            drawings = cell._element.xpath('.//w:drawing')
            if drawings:
                table_drawing_count += len(drawings)
                print(f"✅ 表格 {i} 单元格 [{row_idx}][{cell_idx}] 包含 {len(drawings)} 个 drawing 元素")
                print(f"   文本: {cell.text[:100]}")
                for j, drawing in enumerate(drawings):
                    print(f"   Drawing {j}: {drawing.tag}")

if table_drawing_count == 0:
    print("   ⚠️ 表格中未找到 drawing 元素")

# 3. 检查页眉和页脚
print(f"\n3️⃣ 检查页眉和页脚...\n")

for section_idx, section in enumerate(doc.sections, 1):
    print(f"📄 节 {section_idx}:")

    # 检查页眉
    if section.header:
        print(f"   📋 页眉:")
        print(f"      文本: {section.header.text[:100]}")
        header_drawings = section.header._element.xpath('.//w:drawing')
        if header_drawings:
            print(f"      ✅ 页眉包含 {len(header_drawings)} 个 drawing 元素")
        else:
            print(f"      ⚠️ 页眉无 drawing 元素")

    # 检查页脚
    if section.footer:
        print(f"   📋 页脚:")
        print(f"      文本: {section.footer.text[:100]}")
        footer_drawings = section.footer._element.xpath('.//w:drawing')
        if footer_drawings:
            print(f"      ✅ 页脚包含 {len(footer_drawings)} 个 drawing 元素")
        else:
            print(f"      ⚠️ 页脚无 drawing 元素")

# 4. 检查文档的所有关系
print(f"\n4️⃣ 检查所有图片关系...\n")

if hasattr(doc, 'part') and hasattr(doc.part, 'related_parts'):
    for rId, part in doc.part.related_parts.items():
        content_type = part.content_type
        if 'image' in content_type:
            print(f"✅ 图片 rId: {rId}")
            print(f"   Content-Type: {content_type}")
            print(f"   Blob Size: {len(part.blob)} bytes")

            # 检查这个 rId 在文档中是否被引用
            body_str = str(doc._element.body)
            if rId in body_str:
                print(f"   ✅ 在文档 body 中被引用")
            else:
                print(f"   ⚠️ 未在文档 body 中被引用（可能是删除的图片）")

print(f"\n📊 统计:")
print(f"   - 段落中的 drawing: {drawing_count}")
print(f"   - 表格中的 drawing: {table_drawing_count}")
