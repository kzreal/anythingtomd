"""
Word转换器
将Word文档转换为Markdown，支持切片和图片识别
"""
import logging
import re
from pathlib import Path
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from converters.base import BaseConverter
from utils.zip_helper import ZipHelper
from services.llm_service import ImageRecognitionService
import config


class WordConverter(BaseConverter):
    """Word转Markdown转换器"""

    def __init__(self, file_path: str):
        """
        初始化Word转换器

        Args:
            file_path: Word文件路径
        """
        super().__init__(file_path)
        self.doc: Optional[Document] = None
        self.zip_helper = ZipHelper(Path(file_path).parent.parent / 'downloads')

        # 初始化 LLM 服务
        self.llm_service: Optional[ImageRecognitionService] = None
        if config.LLM_AVAILABLE:
            self.llm_service = ImageRecognitionService(
                endpoint=config.LLM_API_ENDPOINT,
                api_key=config.LLM_API_KEY,
                model=config.LLM_MODEL,
                timeout=config.LLM_TIMEOUT,
                max_retries=config.LLM_MAX_RETRIES
            )

    def load_document(self):
        """加载Word文档"""
        if not self.file_path.exists():
            raise FileNotFoundError(f"文件不存在: {self.file_path}")
        self.doc = Document(str(self.file_path))

    def iter_block_items(self, parent):
        """遍历文档中的所有块元素（段落和表格），保持原始顺序"""
        if hasattr(parent, 'element'):
            parent_elm = parent.element.body
        else:
            parent_elm = parent

        for element in parent_elm.iterchildren():
            if isinstance(element, CT_P):
                yield Paragraph(element, parent)
            elif isinstance(element, CT_Tbl):
                yield Table(element, parent)

    def get_heading_level(self, paragraph: Paragraph) -> int:
        """
        获取段落的标题级别

        Args:
            paragraph: 段落对象

        Returns:
            int: 0=正文, 1=一级标题, 2=二级标题, ...
        """
        # 方法1: 检查样式名称
        if hasattr(paragraph, 'style') and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            if 'heading 1' in style_name or '标题 1' in style_name:
                return 1
            elif 'heading 2' in style_name or '标题 2' in style_name:
                return 2
            elif 'heading 3' in style_name or '标题 3' in style_name:
                return 3
            elif 'heading 4' in style_name or '标题 4' in style_name:
                return 4
            elif 'heading 5' in style_name or '标题 5' in style_name:
                return 5

        # 方法2: 检查大纲级别
        if hasattr(paragraph, '_element'):
            p = paragraph._element
            if p.pPr is not None and hasattr(p.pPr, 'outlineLvl') and p.pPr.outlineLvl is not None:
                return int(p.pPr.outlineLvl.val) + 1

        return 0

    def extract_paragraph_images(self, paragraph: Paragraph) -> List[Dict]:
        """
        检测段落中的图片

        Args:
            paragraph: 段落对象

        Returns:
            图片列表
        """
        images = []

        # 方法1: 检查 wp:inline 格式的图片（内嵌）
        for run in paragraph.runs:
            for inline in run._element.xpath('.//w:drawing/wp:inline'):
                try:
                    blip = inline.xpath('.//a:blip')
                    if blip:
                        embed = inline.xpath('.//a:blip/@r:embed')
                        if embed:
                            # 提取实际图像数据
                            image_data = self._get_image_data(embed[0])
                            if image_data:
                                image_format = self._get_image_format(inline)
                                images.append({
                                    'id': embed[0],
                                    'data': image_data,
                                    'format': image_format,
                                    'placeholder': None
                                })
                except Exception as e:
                    logger.warning(f"Error extracting wp:inline image: {e}")

        # 方法2: 检查 pic:pic 格式的图片（通过 relationship 引用）
        # 这些图片通常在 run 的 text 中显示为 {0}，{0} 或类似占位符
        text = paragraph.text.strip()
        if '{0}' in text or '{1}' in text:
            # 检查所有可能的 pic 引用
            import re
            # 匹配 {0},{0} 或 {0} 或类似模式
            pic_pattern = r'\{[01X]\}|\{[0-9A-F\}|\{[0-9A-F\}'
            if re.search(pic_pattern, text):
                logger.info(f"段落可能包含 pic:pic 图片: {text[:100]}")
                # 查找 pic:nvPic 元素中的图片引用
                for run in paragraph.runs:
                    for drawing in run._element.xpath('.//w:drawing/pic:pic'):
                        for pic_nvPic in drawing.xpath('.//pic:nvPic'):
                            # 获取 pic:blip 引用
                            pic_blip = pic_nvPic.xpath('.//a:blip/@r:embed')
                            if pic_blip:
                                rId = pic_blip[0]
                                # 提取图片数据
                                image_data = self._get_image_data(rId)
                                if image_data:
                                    # 获取图片格式
                                    image_part = self.doc.part.related_parts[rId]
                                    image_format = image_part.content_type.split('/')[-1]
                                    images.append({
                                        'id': rId,
                                        'data': image_data,
                                        'format': image_format,
                                        'placeholder': None
                                    })
                                    logger.info(f"提取到 pic:pic 图片 (ID: {rId}, 大小: {len(image_data)} 字节)")

        return images

    def _get_image_data(self, embed: str) -> Optional[bytes]:
        """从文档关系中提取图像数据"""
        try:
            if hasattr(self.doc, 'part') and hasattr(self.doc.part, 'related_parts'):
                image_part = self.doc.part.related_parts[embed]
                return image_part.blob
        except Exception:
            pass
        return None

    def _get_image_format(self, image_element) -> str:
        """检测图片格式"""
        try:
            blip = image_element.xpath('.//a:blip')
            if blip:
                embed = image_element.xpath('.//a:blip/@r:embed')
                if embed:
                    image_part = self.doc.part.related_parts[embed[0]]
                    return image_part.content_type.split('/')[-1]
        except Exception:
            pass
        return 'png'

    def table_to_markdown_with_images(self, table: Table, start_no: int = 1, processed_images: Dict = None) -> tuple[Optional[str], int]:
        """
        将表格转换为Markdown格式，支持图片合并

        Args:
            table: 表格对象
            start_no: 起始行号
            processed_images: 已处理的图片字典

        Returns:
            (Markdown内容, 下一行号）
        """
        if processed_images is None:
            processed_images = {}

        if not table.rows:
            return None, start_no

        lines = []
        no = start_no

        # 处理表头
        header_cells = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
        lines.append(f"<!-- {no} --> | " + " | ".join(header_cells) + " |")
        lines.append("<!-- " + str(no + 1) + " --> |" + "|".join(["---"] * len(header_cells)) + "|")
        no += 2

        # 处理数据行（包括表头）
        for row in table.rows:
            # 收集当前行的所有图片
            row_images = []
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    row_images.extend(self.extract_paragraph_images(paragraph))

            # 如果该行有图片，合并输出
            if row_images:
                descriptions = []
                for img in row_images:
                    desc = processed_images.get(img['id'])
                    if desc:
                        descriptions.append(f"[图片: {desc}]")
                    else:
                        descriptions.append("[图片: 未识别图片]")

                # 合并图片作为新的一行，使用表格格式
                lines.append(f"<!-- {no} --> | " + " | ".join(descriptions) + " |")
                no += 1

            # 添加表格行内容
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            if not all(cell in ("", " ") for cell in cells):
                lines.append(f"<!-- {no} --> | " + " | ".join(cells) + " |")
                no += 1

        return '\n'.join(lines) + '\n\n', no

    def process_images_batch(self, image_objects: List[Dict]) -> Dict:
        """
        批量处理图片，使用 LLM 识别

        Args:
            image_objects: 图片对象列表

        Returns:
            图片ID到描述的映射字典
        """
        if not config.LLM_AVAILABLE or not self.llm_service:
            logger.warning("LLM service not available, returning None for all images")
            return {img['id']: None for img in image_objects}

        logger.info(f"开始批量处理 {len(image_objects)} 张图片")

        # 过滤需要处理的图片
        images_to_process = []
        for img in image_objects:
            if img['data']:  # 只有有实际数据的图片才处理
                images_to_process.append(img)
            else:
                logger.warning(f"图片 {img['id']} 没有数据，跳过处理")

        if not images_to_process:
            logger.info("No images to process")
            return {img['id']: None for img in image_objects}

        # 限制图片大小（例如 10MB）
        processed_results = {}
        for i, img in enumerate(images_to_process, 1):
            try:
                logger.info(f"正在处理第 {i}/{len(images_to_process)} 张图片: {img['id']}")
                image_size = len(img['data'])
                logger.info(f"图片大小: {image_size / 1024:.2f} KB, 格式: {img['format']}")

                if image_size > 10 * 1024 * 1024:  # 10MB
                    processed_results[img['id']] = None
                    logger.warning(f"Image {img['id']} too large, skipping")
                    continue

                logger.info(f"开始调用 LLM API 识别图片 {img['id']}")
                description = self.llm_service.describe_image(img['data'], img['format'])
                logger.info(f"LLM API 调用完成，结果长度: {len(description) if description else 0}")

                if description:
                    processed_results[img['id']] = description
                    logger.info(f"图片 {img['id']} 处理成功")
                else:
                    processed_results[img['id']] = None
                    logger.warning(f"图片 {img['id']} 识别失败")
            except Exception as e:
                logger.error(f"Error processing image {img['id']}: {e}")
                processed_results[img['id']] = None

        # 为未处理的图片添加 None
        for img in image_objects:
            if img['id'] not in processed_results:
                processed_results[img['id']] = None

        return processed_results

    def table_to_markdown(self, table: Table, start_no: int = 1) -> tuple[Optional[str], int]:
        """
        将表格转换为Markdown格式，带编号

        Args:
            table: 表格对象
            start_no: 起始行号

        Returns:
            (Markdown内容, 下一行号）
        """
        if not table.rows:
            return None, start_no

        lines = []
        no = start_no

        # 表头
        header_cells = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
        lines.append(f"<!-- {no} --> | " + " | ".join(header_cells) + " |")
        no += 1
        lines.append(f"<!-- {no} --> |" + "|".join(["---"] * len(header_cells)) + "|")
        no += 1

        # 数据行
        for row in table.rows[1:]:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            if all(cell in ("", " ") for cell in cells):
                continue
            lines.append(f"<!-- {no} --> | " + " | ".join(cells) + " |")
            no += 1

        return '\n'.join(lines) + '\n\n', no

    def convert(self, options: Dict) -> str:
        """
        转换Word文档为ZIP

        Args:
            options: 转换选项
                - max_level: 切片级别（0/1/2/3/all）

        Returns:
            ZIP文件路径
        """
        self.load_document()

        max_level = options.get('max_level', 0)
        # 处理字符串形式的max_level
        if max_level == 'all':
            max_level = float('inf')
        elif isinstance(max_level, str):
            max_level = int(max_level)

        # 转换文档
        sections = self._slice_document(max_level)
        self.sections = sections

        # 准备输出文件
        files = {}

        # 添加各章节文件
        for section in sections:
            index = section['index'] + 1
            safe_title = self.zip_helper.sanitize_filename(section['title'])
            filename = f"{index:03d}_{safe_title}.md"
            section['filename'] = filename

            content = ''.join(section['content'])
            files[filename] = content

        # 添加索引文件
        index_content = self.zip_helper.create_index(
            original_filename=self.file_path.name,
            file_type='word',
            sections=sections,
            options={'max_level': max_level}
        )
        files['00_index.md'] = index_content

        # 创建ZIP文件
        zip_name = self.get_output_filename()
        zip_path = self.zip_helper.create_zip(files, zip_name)

        return str(zip_path)

    def get_preview_text(self, options: Dict) -> tuple[str, List[Dict]]:
        """
        获取预览文本和所有切片信息

        Args:
            options: 转换选项

        Returns:
            (预览文本, 切片列表)
        """
        self.load_document()

        max_level = options.get('max_level', 0)
        # 处理字符串形式的max_level
        if max_level == 'all':
            max_level = float('inf')
        elif isinstance(max_level, str):
            max_level = int(max_level)

        sections = self._slice_document(max_level)
        self.sections = sections  # 缓存切片信息

        if not sections:
            return "# 空文档", []

        # 返回第一个切片的内容和所有切片信息
        first_section_content = ''.join(sections[0]['content'])

        # 构建切片信息列表（包含完整内容用于缓存）
        sections_info = []
        for section in sections:
            sections_info.append({
                'index': section['index'],
                'title': section['title'],
                'filename': f"{section['index']:03d}_{self.zip_helper.sanitize_filename(section['title'])}.md",
                'content': ''.join(section['content'])
            })

        return first_section_content, sections_info

    def get_sections(self) -> List[Dict]:
        """
        获取所有切片信息（不包含完整内容）

        Returns:
            切片列表，每个切片包含：index, title, filename
        """
        if not hasattr(self, 'sections') or not self.sections:
            # 需要先执行转换
            self.load_document()
            max_level = 0  # 默认值，或根据实际需求调整
            self.sections = self._slice_document(max_level)

        # 返回不含完整内容的切片信息
        return [{
            'index': s['index'],
            'title': s['title'],
            'filename': s.get('filename', f"{s['index']:03d}.md")
        } for s in self.sections]

    def _slice_document(self, max_level):
        """切片文档"""
        if not self.doc:
            self.load_document()

        # 收集所有图片
        all_images = []

        # 遍历文档收集图片
        logger.info(f"开始提取图片，文档包含 {len(self.doc.paragraphs)} 个段落和 {len(self.doc.tables)} 个表格")
        for paragraph in self.doc.paragraphs:
            images = self.extract_paragraph_images(paragraph)
            all_images.extend(images)

        for table in self.doc.tables:
            # 收集表格中的图片
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        images = self.extract_paragraph_images(paragraph)
                        all_images.extend(images)

        logger.info(f"共提取到 {len(all_images)} 张图片")

        # 使用 LLM 处理图片
        processed_images = {}
        if config.LLM_AVAILABLE and self.llm_service and all_images:
            logger.info(f"开始使用 LLM 处理 {len(all_images)} 张图片")
            processed_images = self.process_images_batch(all_images)
            logger.info(f"LLM 图片处理完成，成功处理 {len([k for k, v in processed_images.items() if v])} 张图片")
        else:
            # LLM 不可用时设置为 None，由调用者生成占位符
            logger.warning(f"LLM 不可用或没有图片")
            for img in all_images:
                processed_images[img['id']] = None

        # 零级模式：整个文档为一个文件
        if max_level == 0:
            return [self._create_full_document(processed_images)]

        # 按章节切片
        return self._slice_by_heading(max_level, processed_images)

    def _create_full_document(self, processed_images: Dict = None) -> Dict:
        """创建完整的单一文档"""
        if processed_images is None:
            processed_images = {}

        section = {
            'level': 0,
            'title': self.file_path.stem,
            'content': [],
            'index': 0
        }

        line_no = 1
        for block in self.iter_block_items(self.doc):
            if isinstance(block, Paragraph):
                level = self.get_heading_level(block)
                text = block.text.strip()
                images = self.extract_paragraph_images(block)

                if not text and not images:
                    continue

                # 跳过目录
                if any(kw in text for kw in ['目录', '目  录', 'CONTENTS']):
                    continue

                if level > 0:
                    section['content'].append(f"<!-- {line_no} --> {'#' * level} {text}\n")
                    line_no += 1
                elif text:
                    section['content'].append(f"<!-- {line_no} --> {text}\n")
                    line_no += 1

                # 处理图片
                for img in images:
                    description = processed_images.get(img['id'])
                    if description:
                        # LLM 识别结果
                        section['content'].append(f"<!-- {line_no} -->[图片: {description}]\n")
                    else:
                        # 使用占位符
                        section['content'].append(f"<!-- {line_no} -->[图片: 未识别图片]\n")
                    line_no += 1

            elif isinstance(block, Table):
                table_md, line_no = self.table_to_markdown_with_images(block, line_no, processed_images)
                if table_md:
                    section['content'].append(table_md)

        return section

    def _slice_by_heading(self, max_level, processed_images: Dict = None):
        """按标题切片"""
        if processed_images is None:
            processed_images = {}

        if max_level is None or max_level == 'all':
            max_level = float('inf')

        sections = []
        section_stack = []
        section_index = 0
        line_no = 1

        cover_section = {
            'level': 0,
            'title': '封面',
            'content': [],
            'index': section_index
        }
        section_stack.append(cover_section)

        for block in self.iter_block_items(self.doc):
            if isinstance(block, Paragraph):
                level = self.get_heading_level(block)
                text = block.text.strip()
                images = self.extract_paragraph_images(block)

                if not text and not images:
                    continue

                if any(kw in text for kw in ['目录', '目  录', 'CONTENTS']):
                    continue

                if level > 0 and level <= max_level:
                    if section_stack[-1]['content']:
                        sections.append(section_stack[-1])
                        section_index += 1

                    new_section = {
                        'level': level,
                        'title': text,
                        'content': [],
                        'index': section_index
                    }

                    while section_stack and section_stack[-1]['level'] > level:
                        section_stack.pop()

                    section_stack.append(new_section)
                    section_stack[-1]['content'].append(f"<!-- {line_no} --> {'#' * level} {text}\n")
                    line_no += 1
                else:
                    if level > 0:
                        section_stack[-1]['content'].append(f"<!-- {line_no} --> {'#' * level} {text}\n")
                    elif text:
                        section_stack[-1]['content'].append(f"<!-- {line_no} --> {text}\n")
                    line_no += 1

                # 处理图片
                for img in images:
                    description = processed_images.get(img['id'])
                    if description:
                        # LLM 识别结果
                        section_stack[-1]['content'].append(f"<!-- {line_no} -->[图片: {description}]\n")
                    else:
                        # 使用占位符
                        section_stack[-1]['content'].append(f"<!-- {line_no} -->[图片: 未识别图片]\n")
                    line_no += 1

            elif isinstance(block, Table):
                table_md, line_no = self.table_to_markdown_with_images(block, line_no, processed_images)
                if table_md:
                    section_stack[-1]['content'].append(table_md)

        # 添加剩余章节
        for section in section_stack:
            if section['content']:
                sections.append(section)

        return sections

    def get_output_filename(self) -> str:
        """获取输出文件名"""
        return f"converted_{self.file_path.stem}.zip"

    def cleanup(self):
        """清理资源"""
        if self.llm_service:
            self.llm_service.close()
